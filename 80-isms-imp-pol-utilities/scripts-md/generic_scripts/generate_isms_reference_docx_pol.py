from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_FILE = "isms-reference-pol.docx"

doc = Document()
styles = doc.styles


def add_field_run(paragraph, field_code: str):
    """Insert a Word field code like PAGE or NUMPAGES into a paragraph."""
    run = paragraph.add_run()
    r = run._r
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {field_code} "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r.append(fld_char_begin)
    r.append(instr_text)
    r.append(fld_char_sep)
    r.append(fld_char_end)
    return run


def set_paragraph_style(
    name,
    size,
    bold=False,
    italic=False,
    base="Normal",
    space_before=0,
    space_after=0,
    line_spacing=1,
    font_name="Calibri",
    keep_together=False,
):
    """
    Create or update a paragraph style compatible with Pandoc.
    Compact POL layout (audit-friendly, minimal pages).
    """
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    # Avoid circular reference for base style
    if base != name and base in styles:
        style.base_style = styles[base]

    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = None  # Force black

    pformat = style.paragraph_format
    pformat.space_before = Pt(space_before)
    pformat.space_after = Pt(space_after)
    pformat.line_spacing = line_spacing
    if keep_together:
        pformat.keep_together = True

    return style


# ------------------------------------------------------------------
# CORE PANDOC STYLES - POL (Compact, Audit-Ready)
# ------------------------------------------------------------------

# Base paragraph style
set_paragraph_style(
    "Normal",
    size=10,
    space_before=0,
    space_after=0,
    line_spacing=1,
)

# Title and headings
set_paragraph_style(
    "Title",
    size=18,
    bold=True,
    space_before=0,
    space_after=8,
)

set_paragraph_style(
    "Subtitle",
    size=13,
    italic=True,
    space_before=0,
    space_after=6,
)

set_paragraph_style(
    "Heading 1",
    size=13,
    bold=True,
    space_before=8,
    space_after=4,
    keep_together=True,
)

set_paragraph_style(
    "Heading 2",
    size=11,
    bold=True,
    space_before=6,
    space_after=3,
    keep_together=True,
)

set_paragraph_style(
    "Heading 3",
    size=10,
    bold=True,
    space_before=4,
    space_after=2,
    keep_together=True,
)

set_paragraph_style(
    "Heading 4",
    size=10,
    bold=True,
    space_before=3,
    space_after=2,
)

set_paragraph_style(
    "Heading 5",
    size=10,
    bold=True,
    italic=True,
    space_before=2,
    space_after=2,
)

set_paragraph_style(
    "Heading 6",
    size=10,
    italic=True,
    space_before=2,
    space_after=2,
)

# ------------------------------------------------------------------
# PANDOC-SPECIFIC STYLES (Critical for markdown conversion)
# ------------------------------------------------------------------

# First paragraph after heading (Pandoc sometimes uses this)
set_paragraph_style(
    "First Paragraph",
    size=10,
    base="Normal",
    space_before=0,
    space_after=0,
)

# Body text (alternative to Normal)
set_paragraph_style(
    "Body Text",
    size=10,
    base="Normal",
    space_before=0,
    space_after=0,
)

# Compact paragraph (Pandoc uses for tight spacing)
set_paragraph_style(
    "Compact",
    size=10,
    base="Normal",
    space_before=0,
    space_after=0,
)

# List styles
set_paragraph_style(
    "List Paragraph",
    size=10,
    space_before=0,
    space_after=0,
    line_spacing=1,
)

set_paragraph_style(
    "List Bullet",
    size=10,
    base="List Paragraph",
    space_before=0,
    space_after=0,
)

set_paragraph_style(
    "List Number",
    size=10,
    base="List Paragraph",
    space_before=0,
    space_after=0,
)

# Code blocks (Pandoc uses "Source Code" or "Verbatim")
set_paragraph_style(
    "Source Code",
    size=9,
    font_name="Consolas",
    space_before=4,
    space_after=4,
)

set_paragraph_style(
    "Verbatim",
    size=9,
    font_name="Consolas",
    space_before=4,
    space_after=4,
)

# Block quotes
set_paragraph_style(
    "Block Quote",
    size=10,
    italic=True,
    space_before=4,
    space_after=4,
)

set_paragraph_style(
    "Quote",
    size=10,
    italic=True,
    base="Block Quote",
    space_before=4,
    space_after=4,
)

# Definition lists
set_paragraph_style(
    "Definition Term",
    size=10,
    bold=True,
    space_before=2,
    space_after=0,
)

set_paragraph_style(
    "Definition",
    size=10,
    space_before=0,
    space_after=2,
)

# ------------------------------------------------------------------
# CUSTOM ISMS STYLES
# ------------------------------------------------------------------

set_paragraph_style(
    "Document Metadata",
    size=10,
    bold=True,
    space_after=4,
)

set_paragraph_style(
    "Policy Quote",
    size=10,
    italic=True,
    space_before=4,
    space_after=4,
)

set_paragraph_style(
    "ISMS Table",
    size=9.5,
    space_after=2,
)

# Table styles (Pandoc expects these)
set_paragraph_style(
    "Table Caption",
    size=10,
    italic=True,
    space_before=2,
    space_after=6,
)

set_paragraph_style(
    "Table Heading",
    size=10,
    bold=True,
    space_before=0,
    space_after=0,
)

# Image captions
set_paragraph_style(
    "Image Caption",
    size=9,
    italic=True,
    space_before=2,
    space_after=6,
)

set_paragraph_style(
    "Caption",
    size=9,
    italic=True,
    space_before=2,
    space_after=6,
)

# ------------------------------------------------------------------
# FOOTER WITH DYNAMIC PAGE NUMBERING
# ------------------------------------------------------------------

footer = doc.sections[0].footer
p = footer.paragraphs[0]
p.clear()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add footer content
run = p.add_run("ISMS Policy | Version | Page ")
run.font.name = "Calibri"
run.font.size = Pt(9)
run.font.color.rgb = None

add_field_run(p, "PAGE")

run = p.add_run(" / ")
run.font.name = "Calibri"
run.font.size = Pt(9)
run.font.color.rgb = None

add_field_run(p, "NUMPAGES")

# ------------------------------------------------------------------
# SAVE REFERENCE DOCUMENT
# ------------------------------------------------------------------

doc.save(OUTPUT_FILE)

print(f"✓ Generated POL reference template: {OUTPUT_FILE}")
print(f"✓ Includes {len([s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH])} paragraph styles")
print(f"✓ Pandoc-compatible: Tables, Code Blocks, Lists, Quotes, Headings 1-6")
print(f"✓ Dynamic page numbering: PAGE / NUMPAGES fields")

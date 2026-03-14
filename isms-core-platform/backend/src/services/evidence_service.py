"""Evidence service — file upload, text extraction, DB persistence, OpenSearch indexing.

Supported file types:
  Documents : .docx .pdf .xlsx .txt .md .csv .json
  Images    : .png .jpg .jpeg (stored as-is; no text extraction)
"""

import hashlib
import logging
import re
import uuid
from datetime import date
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.orm import Session as DBSession

from src.core.config import get_settings
from src.database.enums import EvidenceStatus, EvidenceType
from src.domain.compliance import Evidence
from src.services import search_service

logger = logging.getLogger(__name__)

MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB

ALLOWED_EXTENSIONS: set[str] = {
    ".docx", ".pdf", ".xlsx",
    ".png", ".jpg", ".jpeg",
    ".txt", ".md", ".csv", ".json",
}

_IMAGE_EXTS: set[str] = {".png", ".jpg", ".jpeg"}
_TEXT_EXTS: set[str] = {".txt", ".md", ".csv", ".json"}


# ---------------------------------------------------------------------------
# Validation + text extraction
# ---------------------------------------------------------------------------

def validate_extension(filename: str) -> str:
    """Return lowercase extension or raise ValueError."""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"File type '{ext}' not supported. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    return ext


def extract_text(file_path: Path, ext: str) -> str | None:
    """Extract plain text content from a document for OpenSearch indexing.

    Returns None for image files or if extraction fails.
    """
    try:
        if ext == ".docx":
            from docx import Document
            doc = Document(str(file_path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(c.text.strip() for c in row.cells if c.text.strip())
            return "\n".join(parts) or None

        if ext == ".pdf":
            import pdfplumber
            texts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        texts.append(t)
            return "\n\n".join(texts) or None

        if ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"[Sheet: {ws.title}]")
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) for c in row if c is not None)
                    if row_text.strip():
                        parts.append(row_text)
                        row_count += 1
                        if row_count >= 1000:
                            parts.append("... (truncated)")
                            break
            wb.close()
            return "\n".join(parts) or None

        if ext in _TEXT_EXTS:
            return file_path.read_text(encoding="utf-8", errors="replace")[:100_000] or None

        # Images — no text
        return None

    except Exception as e:
        logger.warning("Text extraction failed for %s: %s", file_path.name, e)
        return None


# ---------------------------------------------------------------------------
# File storage
# ---------------------------------------------------------------------------

def _sanitize_filename(stem: str) -> str:
    safe = re.sub(r"[^\w\-.]", "_", stem).strip("_")
    return safe[:80] or "evidence"


def _save_upload(file_bytes: bytes, original_filename: str, group_code: str) -> Path:
    """Save uploaded bytes to disk. group_code='draft' for unassigned files."""
    settings = get_settings()
    ext = Path(original_filename).suffix.lower()
    safe_stem = _sanitize_filename(Path(original_filename).stem)
    uid = uuid.uuid4().hex[:8]
    dest_name = f"{uid}_{safe_stem}{ext}"
    dest_dir = Path(settings.uploads_path) / "evidence" / group_code
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / dest_name
    dest_path.write_bytes(file_bytes)
    return dest_path


# ---------------------------------------------------------------------------
# CRUD operations
# ---------------------------------------------------------------------------

def create_evidence(
    db: DBSession,
    control_group_id: uuid.UUID,
    group_code: str,
    file_bytes: bytes,
    original_filename: str,
    evidence_type: str,
    title: str,
    requirement_id: uuid.UUID | None = None,
    assessment_item_id: uuid.UUID | None = None,
    collected_date: date | None = None,
    expires_date: date | None = None,
    notes: str | None = None,
) -> Evidence:
    """Save file, extract text, create DB record, index to OpenSearch."""
    ext = validate_extension(original_filename)

    if len(file_bytes) > MAX_FILE_BYTES:
        raise ValueError(
            f"File exceeds 50 MB limit ({len(file_bytes) // 1024 // 1024} MB)"
        )

    dest_path = _save_upload(file_bytes, original_filename, group_code)

    extracted_text = extract_text(dest_path, ext)
    word_count = len(extracted_text.split()) if extracted_text else 0
    content_hash = hashlib.sha256(file_bytes).hexdigest()

    evidence = Evidence(
        id=uuid.uuid4(),
        control_group_id=control_group_id,
        requirement_id=requirement_id,
        assessment_item_id=assessment_item_id,
        evidence_type=EvidenceType(evidence_type),
        evidence_status=EvidenceStatus.ACTIVE,
        title=title,
        file_path=str(dest_path),
        collected_date=collected_date,
        expires_date=expires_date,
        metadata_={
            "original_filename": original_filename,
            "file_size": len(file_bytes),
            "file_ext": ext,
            "content_hash": content_hash,
            "word_count": word_count,
            "notes": notes or "",
            "text_preview": (extracted_text or "")[:500],
        },
    )
    db.add(evidence)
    db.flush()

    # Index to OpenSearch — graceful: won't fail upload if OS is down
    if extracted_text:
        search_service.index_evidence(
            evidence_id=str(evidence.id),
            title=title,
            evidence_type=evidence_type,
            control_group_code=group_code,
            original_filename=original_filename,
            text=extracted_text,
            word_count=word_count,
            file_path=str(dest_path),
        )

    return evidence


def create_draft_evidence(
    db: DBSession,
    file_bytes: bytes,
    original_filename: str,
    evidence_type: str,
    title: str,
    notes: str | None = None,
) -> Evidence:
    """Upload a file as a draft — no control group required yet."""
    ext = validate_extension(original_filename)

    if len(file_bytes) > MAX_FILE_BYTES:
        raise ValueError(
            f"File exceeds 50 MB limit ({len(file_bytes) // 1024 // 1024} MB)"
        )

    dest_path = _save_upload(file_bytes, original_filename, "draft")
    extracted_text = extract_text(dest_path, ext)
    word_count = len(extracted_text.split()) if extracted_text else 0
    content_hash = hashlib.sha256(file_bytes).hexdigest()

    evidence = Evidence(
        id=uuid.uuid4(),
        control_group_id=None,
        evidence_type=EvidenceType(evidence_type),
        evidence_status=EvidenceStatus.DRAFT,
        title=title,
        file_path=str(dest_path),
        metadata_={
            "original_filename": original_filename,
            "file_size": len(file_bytes),
            "file_ext": ext,
            "content_hash": content_hash,
            "word_count": word_count,
            "notes": notes or "",
            "text_preview": (extracted_text or "")[:500],
        },
    )
    db.add(evidence)
    db.flush()
    return evidence


def get_evidence(db: DBSession, evidence_id: uuid.UUID) -> Evidence | None:
    return db.get(Evidence, evidence_id)


def list_evidence(
    db: DBSession,
    control_group_id: uuid.UUID | None = None,
    evidence_type: str | None = None,
    evidence_status: str | None = None,
    limit: int = 100,
    offset: int = 0,
) -> list[Evidence]:
    stmt = select(Evidence).order_by(Evidence.created_at.desc())
    if control_group_id:
        stmt = stmt.where(Evidence.control_group_id == control_group_id)
    if evidence_type:
        stmt = stmt.where(Evidence.evidence_type == EvidenceType(evidence_type))
    if evidence_status:
        stmt = stmt.where(Evidence.evidence_status == EvidenceStatus(evidence_status))
    return db.execute(stmt.limit(limit).offset(offset)).scalars().all()


def update_evidence(db: DBSession, evidence: Evidence, updates: dict) -> Evidence:
    """Apply partial updates to an evidence record."""
    if "title" in updates and updates["title"]:
        evidence.title = updates["title"]
    if "evidence_type" in updates and updates["evidence_type"]:
        evidence.evidence_type = EvidenceType(updates["evidence_type"])
    if "collected_date" in updates:
        evidence.collected_date = updates["collected_date"]
    if "expires_date" in updates:
        evidence.expires_date = updates["expires_date"]
    if "verified_by" in updates:
        evidence.verified_by = updates["verified_by"]
    if "verified_date" in updates:
        evidence.verified_date = updates["verified_date"]
    if "notes" in updates and updates["notes"] is not None:
        # Reassign JSONB dict to trigger SQLAlchemy dirty tracking
        meta = dict(evidence.metadata_)
        meta["notes"] = updates["notes"]
        evidence.metadata_ = meta
    db.flush()
    return evidence


def delete_evidence(db: DBSession, evidence: Evidence) -> None:
    """Delete DB record and uploaded file."""
    file_path = Path(evidence.file_path)
    evidence_id = str(evidence.id)
    db.delete(evidence)
    db.flush()
    try:
        if file_path.exists():
            file_path.unlink()
    except OSError as e:
        logger.warning("Could not delete evidence file %s: %s", file_path, e)
    # Remove from OpenSearch (graceful)
    search_service.delete_evidence(evidence_id)

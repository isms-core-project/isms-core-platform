"""Word (.docx) policy parser — extracts structured content from Word documents."""

import logging
import re
from pathlib import Path

from src.importers.parsers.base import (
    BasePolicyParser,
    ParsedPolicy,
    Section,
    compute_content_hash,
    derive_group_code,
    detect_product_and_type,
)

logger = logging.getLogger(__name__)

_DOC_ID_RE = re.compile(r"(ISMS-(?:OP-)?POL-[A-Za-z0-9][\w.\-]+)")

# Heading style prefixes used by python-docx
_HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3, "Heading 4": 4}


class DocxPolicyParser(BasePolicyParser):
    """Parses policy documents from Word .docx files using python-docx."""

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> ParsedPolicy:
        from docx import Document

        content_hash = compute_content_hash(file_path)
        doc = Document(str(file_path))

        # Extract all paragraphs with their style info
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            level = _HEADING_STYLES.get(style_name, 0)
            paragraphs.append((text, level, style_name))

        # Extract table content (for Document Control)
        table_metadata = self._extract_table_metadata(doc)

        # Build full text for word count and ID extraction
        full_text = "\n".join(t for t, _, _ in paragraphs)
        if not full_text.strip():
            raise ValueError(f"No text extracted from DOCX: {file_path}")

        document_id = self._extract_document_id(
            full_text, table_metadata, file_path
        )
        product_type, policy_type = detect_product_and_type(document_id)
        group_code = derive_group_code(document_id)
        title = self._extract_title(paragraphs, table_metadata, document_id)
        sections = self._extract_sections(paragraphs)
        word_count = len(full_text.split())

        metadata = table_metadata.copy()
        metadata["source_format"] = "docx"

        return ParsedPolicy(
            document_id=document_id,
            title=title,
            product_type=product_type,
            policy_type=policy_type,
            group_code=group_code,
            content_hash=content_hash,
            word_count=word_count,
            sections=sections,
            metadata=metadata,
        )

    def _extract_table_metadata(self, doc) -> dict:
        """Extract Document Control table fields."""
        metadata = {}
        field_map = {
            "document id": "document_id",
            "document title": "title",
            "document type": "document_type",
            "version": "version",
            "classification": "classification",
            "status": "status",
            "document owner": "owner",
            "approved by": "approved_by",
        }

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    field_name = cells[0].lower().strip("*")
                    key = field_map.get(field_name)
                    if key:
                        metadata[key] = cells[1]
        return metadata

    def _extract_document_id(
        self, full_text: str, table_metadata: dict, file_path: Path
    ) -> str:
        """Find ISMS document ID from text, table, or filename."""
        # Check table metadata first
        if "document_id" in table_metadata:
            doc_id = table_metadata["document_id"]
            m = _DOC_ID_RE.search(doc_id)
            if m:
                return m.group(1).rstrip(" \t—-–:*")

        # Search text
        m = _DOC_ID_RE.search(full_text[:1000])
        if m:
            return m.group(1).rstrip(" \t—-–:*")

        # Filename fallback
        m = _DOC_ID_RE.search(file_path.stem)
        if m:
            return m.group(1).rstrip(" \t—-–:*")

        raise ValueError(f"Cannot extract document_id from DOCX: {file_path}")

    def _extract_title(
        self, paragraphs: list[tuple], table_metadata: dict, document_id: str
    ) -> str:
        """Extract title from table metadata, first heading, or document ID."""
        if "title" in table_metadata and table_metadata["title"]:
            return table_metadata["title"]

        # First heading
        for text, level, _ in paragraphs:
            if level > 0:
                return text
            # First bold-looking line with the doc ID
            if document_id.split("-")[-1] in text:
                for sep in ["—", "–", "-", ":"]:
                    if sep in text:
                        title = text.split(sep, 1)[1].strip()
                        if title:
                            return title
        return document_id

    def _extract_sections(self, paragraphs: list[tuple]) -> list[Section]:
        """Group paragraphs into heading + body sections."""
        sections: list[Section] = []
        current_heading = "(preamble)"
        current_level = 0
        current_body_parts: list[str] = []

        for text, level, _ in paragraphs:
            if level > 0:
                # Flush previous section
                if current_body_parts:
                    sections.append(
                        Section(
                            heading=current_heading,
                            body="\n".join(current_body_parts),
                            level=current_level,
                        )
                    )
                current_heading = text
                current_level = level
                current_body_parts = []
            else:
                current_body_parts.append(text)

        # Flush last section
        if current_body_parts:
            sections.append(
                Section(
                    heading=current_heading,
                    body="\n".join(current_body_parts),
                    level=current_level,
                )
            )

        return sections
